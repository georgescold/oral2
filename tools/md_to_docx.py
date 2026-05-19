"""
Convertit un fichier Markdown CAPEPS Oral 2 en .docx propre (style Loys).
Usage : python md_to_docx.py <input.md> [output.docx]
Si output.docx n'est pas fourni, le .docx est créé à côté du .md avec le même nom.
"""
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    bullet_re = re.compile(r"^\s*[-*+]\s+(.*)$")
    numbered_re = re.compile(r"^\s*\d+\.\s+(.*)$")
    heading_re = re.compile(r"^(#{1,6})\s+(.*)$")
    bold_re = re.compile(r"\*\*(.+?)\*\*")
    italic_re = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")

    def add_runs(paragraph, line: str) -> None:
        pos = 0
        tokens = []
        for m in bold_re.finditer(line):
            if m.start() > pos:
                tokens.append(("plain", line[pos:m.start()]))
            tokens.append(("bold", m.group(1)))
            pos = m.end()
        if pos < len(line):
            tokens.append(("plain", line[pos:]))

        final_tokens = []
        for kind, value in tokens:
            if kind == "bold":
                final_tokens.append((kind, value))
                continue
            sub_pos = 0
            for m in italic_re.finditer(value):
                if m.start() > sub_pos:
                    final_tokens.append(("plain", value[sub_pos:m.start()]))
                final_tokens.append(("italic", m.group(1)))
                sub_pos = m.end()
            if sub_pos < len(value):
                final_tokens.append(("plain", value[sub_pos:]))

        for kind, value in final_tokens:
            run = paragraph.add_run(value)
            if kind == "bold":
                run.bold = True
            elif kind == "italic":
                run.italic = True

    in_code = False
    for raw in text.splitlines():
        line = raw.rstrip()

        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph()
            p.add_run(line).font.name = "Consolas"
            continue

        if not line.strip():
            doc.add_paragraph()
            continue

        h = heading_re.match(line)
        if h:
            level = min(len(h.group(1)), 4)
            doc.add_heading(h.group(2).strip(), level=level)
            continue

        b = bullet_re.match(line)
        if b:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, b.group(1).strip())
            continue

        n = numbered_re.match(line)
        if n:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, n.group(1).strip())
            continue

        if line.strip().startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_runs(p, line.strip()[1:].lstrip())
            continue

        p = doc.add_paragraph()
        add_runs(p, line)

    doc.save(str(docx_path))


def main() -> int:
    if len(sys.argv) < 2:
        print("Usage: python md_to_docx.py <input.md> [output.docx]", file=sys.stderr)
        return 1
    md_path = Path(sys.argv[1]).resolve()
    if not md_path.exists():
        print(f"[ERREUR] Fichier introuvable : {md_path}", file=sys.stderr)
        return 2
    docx_path = (
        Path(sys.argv[2]).resolve()
        if len(sys.argv) >= 3
        else md_path.with_suffix(".docx")
    )
    md_to_docx(md_path, docx_path)
    print(f"[OK] {docx_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
